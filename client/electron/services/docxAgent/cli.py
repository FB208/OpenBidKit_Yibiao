# -*- coding: utf-8 -*-
"""docx-agent JSON CLI.

stdin 一条 JSON 请求，stdout 一条 JSON 响应；失败退出码非 0 且 ok:false。
封装 client/vendor/docx-agent 的高保真 docx 操作，供 Electron Main 与 Agent 工具调用。
允许的 op：inspect / locate / extract_range / copy_range。

注意：clipboard 的 copy_* 函数内部各自重新打开 src/tgt 文档，锚点只能传
body 段落索引（int），不能传另一个文档实例里的元素引用。
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

_VENDOR_DIR = Path(__file__).resolve().parents[3] / "vendor" / "docx-agent"
if str(_VENDOR_DIR) not in sys.path:
    sys.path.insert(0, str(_VENDOR_DIR))

from docx import Document  # noqa: E402

from src import clipboard, deleter, locator, reader  # noqa: E402

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_MAX_ALTERNATIVES = 10
_PREVIEW_MAX_CHARS = 1200


class OpError(Exception):
    pass


def _w(name):
    return _W_NS + name


def _load_doc(path):
    if not path:
        raise OpError("缺少 doc 路径")
    return Document(path)


def _style_map(doc):
    return locator._style_id_to_name(doc if hasattr(doc, "part") else None)


def _paragraphs(doc):
    return locator._body_paragraphs(doc)


def _paragraph_text(p):
    return locator._paragraph_text(p)


def _heading_level(doc, p, style_map=None):
    return locator._heading_level_of(p, style_map or _style_map(doc))


def _next_heading_end(doc, start_index, start_level):
    """返回下一个同级或更高级标题的段落索引（end_index_exclusive），没有则 None。"""
    style_map = _style_map(doc)
    paras = _paragraphs(doc)
    for i in range(start_index + 1, len(paras)):
        level = _heading_level(doc, paras[i], style_map)
        if level is not None and (start_level is None or level <= start_level):
            return i
    return None


def _has_parent_heading(doc, cand_index, cand_level, parent_text):
    """cand 之前是否存在层级更浅（level < cand_level）且文本含 parent_text 的标题。"""
    style_map = _style_map(doc)
    paras = _paragraphs(doc)
    for i in range(cand_index - 1, -1, -1):
        level = _heading_level(doc, paras[i], style_map)
        if level is None or level >= cand_level:
            continue
        if parent_text in _paragraph_text(paras[i]):
            return True
    return False


def _collect_hits(fn, doc, **kwargs):
    """按 occurrence 递增调用直到失败，收集候选（封顶判断多命中）。"""
    hits = []
    for occurrence in range(1, _MAX_ALTERNATIVES + 2):
        try:
            hits.append(fn(doc, occurrence=occurrence, **kwargs))
        except (locator.LocateError, ValueError):
            break
    return hits


def _alternatives_from_hits(hits):
    alts = []
    for h in hits[1:]:
        alts.append({
            "paragraph_index": h.get("paragraph_index"),
            "table_index": h.get("table_index"),
            "text": h.get("text") or " / ".join(h.get("header_cells") or []),
        })
    return alts[:_MAX_ALTERNATIVES]


def _paragraph_match_result(doc, hit, strategy, match_count, alternatives):
    index = hit["paragraph_index"]
    level = _heading_level(doc, hit["p_elem"])
    return {
        "strategy": strategy,
        "start_kind": "paragraph",
        "paragraph_index": index,
        "table_index": None,
        "end_index_exclusive": _next_heading_end(doc, index, level),
        "text": hit.get("text", ""),
        "heading_level": level,
        "match_count": match_count,
        "alternatives": alternatives,
    }


def _op_locate(req):
    doc = _load_doc(req.get("doc"))
    strategy = req.get("strategy") or "heading"
    text = (req.get("text") or "").strip()
    section_text = (req.get("section_text") or "").strip()
    keywords = [str(k).strip() for k in (req.get("keywords") or []) if str(k).strip()]

    if strategy == "heading":
        if not text:
            raise OpError("strategy=heading 需要 text")
        level = req.get("level")
        levels = [int(level)] if level else [1, 2, 3, 4, 5]
        for lv in levels:
            hits = _collect_hits(locator.locate_by_heading, doc, text=text, level=lv)
            if not hits:
                continue
            chosen = hits[0]
            if len(hits) > 1 and section_text:
                for cand in hits:
                    if _has_parent_heading(doc, cand["paragraph_index"], lv, section_text):
                        chosen = cand
                        break
            return _paragraph_match_result(
                doc, chosen, "heading", len(hits), _alternatives_from_hits(hits))
        raise OpError(f"未找到标题匹配: {text!r}")

    if strategy == "section":
        if not (text and section_text):
            raise OpError("strategy=section 需要 text 与 section_text")
        level = int(req.get("level") or 1)
        hit = locator.locate_in_section(doc, section_text, text, level=level)
        return _paragraph_match_result(doc, hit, "section", hit.get("match_count", 1), [])

    if strategy == "text":
        if not text:
            raise OpError("strategy=text 需要 text")
        hits = _collect_hits(locator.locate_by_text, doc, text=text)
        if not hits:
            raise OpError(f"未找到文本匹配: {text!r}")
        chosen = hits[0]
        if len(hits) > 1 and section_text:
            for cand in hits:
                cand_level = _heading_level(doc, cand["p_elem"])
                if _has_parent_heading(doc, cand["paragraph_index"],
                                       cand_level or 99, section_text):
                    chosen = cand
                    break
        return _paragraph_match_result(doc, chosen, "text", len(hits), _alternatives_from_hits(hits))

    if strategy == "table_header":
        if not keywords:
            raise OpError("strategy=table_header 需要 keywords")
        hits = _collect_hits(locator.locate_table_by_header, doc, header_keywords=keywords)
        if not hits:
            raise OpError(f"未找到表头匹配: {keywords!r}")
        chosen = hits[0]
        return {
            "strategy": "table_header",
            "start_kind": "table",
            "paragraph_index": None,
            "table_index": chosen["table_index"],
            "end_index_exclusive": None,
            "text": " / ".join(chosen.get("header_cells") or []),
            "heading_level": None,
            "match_count": len(hits),
            "alternatives": _alternatives_from_hits(hits),
        }

    raise OpError(f"未知 strategy: {strategy!r}")


def _table_text(tbl):
    parts = []
    for tr in tbl.findall(_w("tr")):
        cells = [reader._tc_text(tc).strip() for tc in tr.findall(_w("tc"))]
        parts.append(" | ".join(c for c in cells if c))
    return "\n".join(parts)


def _range_preview(doc, start_index, end_index_exclusive):
    """范围内段落与表格拼接的纯文本预览，供正文页展示。"""
    body = locator._body(doc)
    children = [c for c in body if not c.tag.endswith("}sectPr")]
    start_bi = children.index(_paragraphs(doc)[start_index])
    if end_index_exclusive is None:
        end_bi = len(children)
    else:
        end_bi = children.index(_paragraphs(doc)[end_index_exclusive])
    parts = []
    for child in children[start_bi:end_bi]:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            t = _paragraph_text(child).strip()
            if t:
                parts.append(t)
        elif tag == "tbl":
            t = _table_text(child)
            if t:
                parts.append(t)
    preview = "\n".join(parts)
    if len(preview) > _PREVIEW_MAX_CHARS:
        preview = preview[:_PREVIEW_MAX_CHARS] + "…"
    return preview


def _op_extract_range(req):
    src = req.get("src")
    out = req.get("out")
    if not (src and out):
        raise OpError("extract_range 需要 src 与 out")
    start_kind = req.get("start_kind") or "paragraph"

    # 先写一个带空锚点段的文档作为粘贴目标（内容插在锚点段之后）
    tgt_seed = Document()
    if not _paragraphs(tgt_seed):
        tgt_seed.add_paragraph("")
    tgt_seed.save(out)

    if start_kind == "table":
        table_index = int(req.get("table_index", -1))
        doc = _load_doc(src)
        tables = [c for c in doc.element.body if c.tag.endswith("}tbl")]
        if table_index < 0 or table_index >= len(tables):
            raise OpError(f"表索引越界: {table_index}（共 {len(tables)} 张表）")
        result = clipboard.copy_table_across_docs(src, table_index, out, 0, output_path=out)
        preview = _table_text(tables[table_index])
        if len(preview) > _PREVIEW_MAX_CHARS:
            preview = preview[:_PREVIEW_MAX_CHARS] + "…"
    else:
        start = int(req.get("start", -1))
        end = req.get("end")
        expect_text = req.get("expect_text")
        doc = _load_doc(src)
        locator.locate_by_paragraph_index(doc, start, expect_text_contains=expect_text)
        end_i = int(end) if end is not None else None
        if end_i is not None and end_i <= start:
            raise OpError(f"end({end_i}) 必须大于 start({start})")
        result = clipboard.copy_range_across_docs(
            src, start, end_i, out, 0, output_path=out)
        preview = _range_preview(doc, start, end_i)

    return {
        "out": out,
        "inserted_count": result.get("inserted_count"),
        "copied_images": result.get("copied_images", 0),
        "preview_text": preview,
    }


def _snapshot_content_start_index(src_doc):
    """快照结构：[空锚点段, 内容...]。返回首个内容元素的描述。"""
    children = [c for c in src_doc.element.body if not c.tag.endswith("}sectPr")]
    if len(children) < 2:
        raise OpError("快照内容为空")
    first = children[1]
    if first.tag.endswith("}tbl"):
        return {"kind": "table", "table_index": 0}
    return {"kind": "paragraph", "paragraph_index": _paragraphs(src_doc).index(first)}


def _op_copy_range(req):
    src = req.get("src")
    tgt = req.get("tgt")
    out = req.get("out") or tgt
    anchor_bookmark = req.get("anchor_bookmark")
    if not (src and tgt):
        raise OpError("copy_range 需要 src 与 tgt")
    if not anchor_bookmark:
        raise OpError("copy_range 需要 anchor_bookmark")

    # 书签段落在目标文档中的段落索引（copy_* 内部按索引重新解析锚点）
    tgt_probe = _load_doc(tgt)
    bookmark_hit = locator.locate_by_bookmark(tgt_probe, anchor_bookmark)
    tgt_anchor_index = bookmark_hit["paragraph_index"]

    start = req.get("start")
    if start is None:
        # 自动模式：粘贴快照全部内容（跳过锚点空段）
        src_doc = _load_doc(src)
        content = _snapshot_content_start_index(src_doc)
        if content["kind"] == "table":
            result = clipboard.copy_table_across_docs(
                src, content["table_index"], tgt, tgt_anchor_index, output_path=out)
        else:
            result = clipboard.copy_range_across_docs(
                src, content["paragraph_index"], None, tgt, tgt_anchor_index,
                output_path=out)
    else:
        start_i = int(start)
        end_i = int(req["end"]) if req.get("end") is not None else None
        result = clipboard.copy_range_across_docs(
            src, start_i, end_i, tgt, tgt_anchor_index, output_path=out)

    # 粘贴成功后删除书签所在占位段
    tgt_doc = _load_doc(out)
    deleter.delete_paragraph(tgt_doc, locator.locate_by_bookmark(tgt_doc, anchor_bookmark))
    tgt_doc.save(out)

    return {
        "out": out,
        "inserted_count": result.get("inserted_count"),
        "copied_images": result.get("copied_images", 0),
    }


def _op_inspect(req):
    doc = _load_doc(req.get("doc"))
    style_map = _style_map(doc)
    paras = _paragraphs(doc)
    tables = [c for c in doc.element.body if c.tag.endswith("}tbl")]
    headings = []
    for i, p in enumerate(paras):
        level = _heading_level(doc, p, style_map)
        if level is not None and len(headings) < 500:
            headings.append({"index": i, "level": level, "text": _paragraph_text(p)})
    return {"paragraphs": len(paras), "tables": len(tables), "headings": headings}


_OPS = {
    "inspect": _op_inspect,
    "locate": _op_locate,
    "extract_range": _op_extract_range,
    "copy_range": _op_copy_range,
}


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stdin.reconfigure(encoding="utf-8")
    except Exception:
        pass
    try:
        req = json.loads(sys.stdin.read() or "{}")
        op = req.get("op")
        handler = _OPS.get(op)
        if not handler:
            raise OpError(f"未知 op: {op!r}（允许: {', '.join(_OPS)}）")
        result = handler(req)
        print(json.dumps({"ok": True, **(result or {})}, ensure_ascii=False))
    except Exception as exc:  # noqa: BLE001 统一转 JSON 错误
        print(json.dumps({"ok": False, "error": f"{type(exc).__name__}: {exc}"},
                         ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
